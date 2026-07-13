from docx import Document
import os


def export_summary_to_docx(summary, output_folder="exports"):

    os.makedirs(output_folder, exist_ok=True)

    file_path = os.path.join(output_folder, "AI_Summary.docx")

    document = Document()

    document.add_heading("AI Notes Summary", level=1)

    document.add_paragraph(summary)

    document.save(file_path)

    return file_path